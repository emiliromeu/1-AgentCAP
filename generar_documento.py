"""
Genera el document .docx del CAP de mercaderies.
Disseny: Calibri, color de marca #2E7D8F, capçaleres tintades, bordes suaus, logo al header.
"""

from datetime import date, datetime
from docx import Document
from docx.shared import Inches, Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pathlib import Path

from ensamblaje import cronograma_por_semanas
from datos.plantilla_mercancias import AMPLIACION_MERCANCIAS
from datos.plantilla_viatgers   import AMPLIACION_VIATGERS
from herramientas.motor_practicas import colocar_practicas
from herramientas.motor_horario   import construir_dias_lectivos, construir_franjas_semanales

_BASE = Path(__file__).resolve().parent

# ── Constants generals ─────────────────────────────────────────────────────────

DIES_SETMANA_CA = {
    0: "DILLUNS",
    1: "DIMARTS",
    2: "DIMECRES",
    3: "DIJOUS",
    4: "DIVENDRES",
    5: "DISSABTE",
}

DADES_CURS = {
    "formacio":     "Mòduls 1, 2, 3 i 5",
    "durada":       "accelerat 130 hores",
    "empresa":      "AUTOESCOLA OLIVELLA",
    "nif":          "B-60723152",
    "autoritzacio": "CAP-0151",
}

# Amplades de columna en DXA (1 polzada = 1440 DXA)
_W_DIA   = 1400
_W_HORA  = 1800
_W_TEMA  = 4200
_W_PROF  = 3000
_W_TOTAL = _W_DIA + _W_HORA + _W_TEMA + _W_PROF   # 10400 DXA

RUTA_LOGO = str(_BASE / "assets" / "logo_olivella.png")

# ── Colors i tipografia de marca ──────────────────────────────────────────────

_FONT            = 'Calibri'
_COLOR_BRAND     = "2E7D8F"                    # teal del logo Olivella
_COLOR_BRAND_RGB = RGBColor(0x2E, 0x7D, 0x8F)
_COLOR_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
_COLOR_TOTAL     = "EEEEEE"   # gris clar per a files de total
_COLOR_BORDER    = "B0C4C8"   # gris-blavós suau per als bordes
_COLOR_RED_RGB   = RGBColor(0xFF, 0x00, 0x00)  # vermell ampliació cronograma

_AMPLIACION_CURS = {
    "mercancias": AMPLIACION_MERCANCIAS,
    "viatgers":   AMPLIACION_VIATGERS,
}

# Nom del curs segons tipo_curso — usat al títol de portada i als textos
# "QUALIFICACIÓ INICIAL ...". Abans estava hardcodejat a "MERCADERIES" sempre,
# fins i tot en documents de viatgers.
_NOM_CURS = {
    "mercancias": "MERCADERIES",
    "viatgers":   "VIATGERS",
}

# Text llarg de la línia "CURS:" de portada, només per a viatgers (mercaderies
# manté "QUALIFICACIÓ INICIAL MERCADERIES" sense canvis). Petició explícita:
# aquest text NO afecta el títol gran de portada ni el text de COMUNICACIONS,
# que segueixen usant _NOM_CURS ("VIATGERS") tal com estaven.
_TEXT_CURS_PORTADA_VIATGERS = (
    "FORM. MATERIA DE SENSIBIL. VIATGERS AMB DISCAPAC / FORM. TRANSPORTE "
    "ESCOLAR, PRIMERS AUXILIS I TACOGRAF DIG."
)

# ── Cronograma: materials i categories ───────────────────────────────────────

MATERIAS_CRONOGRAMA = [
    ("1.1",     "Cadena cinem."),
    ("1.2",     "Disp. seg."),
    ("1.3",     "Optim. consum"),
    ("1.3 bis", "Anticip. risc"),
    ("1.4",     "Càrrega/desc."),
    ("2.1",     "Entorn social"),
    ("2.2",     "Reglam. merc."),
    ("3.1",     "Riscos/acc."),
    ("3.2",     "Prev. delinq."),
    ("3.3",     "Prev. fís."),
    ("3.4",     "Aptitud"),
    ("3.5",     "Avaluar emerg."),
    ("3.6",     "Imatge marca"),
    ("3.7",     "Entorn econ."),
    ("MM.PP",   "Merc. perill."),
    ("Mòdul 2", "Cisternes"),
]

CATEGORIES_CRONOGRAMA = [
    ("CONDUCCIÓ RACIONAL BASADA EN NORMES DE SEGURETAT",         0,  4),
    ("APLICACIÓ DE LA REGLAMENTACIÓ",                             5,  6),
    ("Salut, seguretat viària i mediambiental, servei logística", 7, 13),
    ("FORMACIÓ COMPLEMENTARIA MERCADERIES",                      14, 15),
]

MATERIAS_CRONOGRAMA_VIATGERS = [
    ("1.1",       "Cadena cinem."),
    ("1.2",       "Disp. seg."),
    ("1.3",       "Optim. consum"),
    ("1.3 bis",   "Anticip. risc"),
    ("1.5",       "Seg. comoditat"),
    ("1.6",       "Op. càrrega"),
    ("2.1",       "Entorn social"),
    ("2.3",       "Reglam. viatg."),
    ("3.1",       "Riscos"),
    ("3.2",       "Prev. delinq."),
    ("3.3",       "Riscos fís."),
    ("3.4",       "Aptitud"),
    ("3.5",       "Sit. emerg."),
    ("3.6",       "Imatge marca"),
    ("3.8",       "Entorn econ. v."),
    ("Mòdul 1",   "Sensib. discap."),
    ("Mòdul 2",   "Transport Esc."),
    ("Mòdul 3",   "Primers Aux."),
    ("Mòdul 5",   "Tacògraf Dig."),
]

CATEGORIES_CRONOGRAMA_VIATGERS = [
    ("CONDUCCIÓ RACIONAL BASADA EN NORMES DE SEGURETAT",             0,  5),
    ("APLICACIÓ DE LA REGLAMENTACIÓ",                                 6,  7),
    ("SALUT, SEGURETAT VIÀRIA I MEDIAMBIENTAL, SERVEI, LOGÍSTICA",    8, 14),
    ("FORMACIÓ COMPLEMENTÀRIA VIATGERS",                             15, 18),
]

_MATERIAS_CURS = {
    "mercancias": (MATERIAS_CRONOGRAMA,         CATEGORIES_CRONOGRAMA),
    "viatgers":   (MATERIAS_CRONOGRAMA_VIATGERS, CATEGORIES_CRONOGRAMA_VIATGERS),
}

# Amplades per al cronograma (pàgina apaïsada, marges 0,25")
# Total disponible: 16838 - 2×360 = 16118 Twips
_WC_SETMANA = 1300
_WC_MATERIA  = 860     # 16 columnes × 860 = 13760
_WC_TOTAL    = 1058    # 1300 + 13760 + 1058 = 16118 ✓


# ── Format d'hores ────────────────────────────────────────────────────────────

def _fmt_hora(t):
    """Temps en format CAP: '18' per hores exactes, '20'15' per minuts."""
    if t.minute == 0:
        return str(t.hour)
    return f"{t.hour}'{t.minute:02d}"

def _hora_rang(ini, fi):
    return f"{_fmt_hora(ini)} a {_fmt_hora(fi)}"

def _total_hores_str(tramos):
    base = date.today()
    total_min = sum(
        int(
            (datetime.combine(base, t["fin"]) - datetime.combine(base, t["inicio"]))
            .total_seconds() / 60
        )
        for t in tramos if t["tipo"] == "clase"
    )
    h, m = total_min // 60, total_min % 60
    return f"{h} hores" if m == 0 else f"{h}'{m:02d} hores"

def _fmt_hores_cron(h):
    return f"{int(h)} h" if h == int(h) else f"{h:g} h"

def _fmt_hora_practica(hora_str):
    h, m = hora_str.split(':')
    return f"{int(h)}:{m}h"


# ── Helpers XML ───────────────────────────────────────────────────────────────

def _set_cell_width(cell, w_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcW')):
        tcPr.remove(ex)
    el = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(w_dxa))
    el.set(qn('w:type'), 'dxa')
    tcPr.insert(0, el)

def _set_table_width(table, w_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for ex in tblPr.findall(qn('w:tblW')):
        tblPr.remove(ex)
    el = OxmlElement('w:tblW')
    el.set(qn('w:w'), str(w_dxa))
    el.set(qn('w:type'), 'dxa')
    tblPr.append(el)

def _set_tbl_grid(table, col_widths):
    tbl = table._tbl
    existing = tbl.find(qn('w:tblGrid'))
    if existing is not None:
        tbl.remove(existing)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tbl_pr = tbl.find(qn('w:tblPr'))
    tbl.insert(list(tbl).index(tbl_pr) + 1, tblGrid)

def _set_table_cell_margin(table, top=60, right=80, bottom=60, left=80):
    """Marges interiors per defecte de totes les cel·les (1pt = 20dxa)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for ex in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(ex)
    tcm = OxmlElement('w:tblCellMar')
    for side, val in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcm.append(el)
    tblPr.append(tcm)

def _add_borders(cell, sz=4, color=None):
    """Vores fines. Per defecte usa el gris-blavós suau de marca."""
    if color is None:
        color = _COLOR_BORDER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(ex)
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def _set_shading(cell, fill=_COLOR_TOTAL):
    """Sombreig de cel·la. val='clear' ≡ ShadingType.CLEAR — NO usar 'solid'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:shd')):
        tcPr.remove(ex)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _set_vmerge_start(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:vMerge')):
        tcPr.remove(ex)
    vm = OxmlElement('w:vMerge')
    vm.set(qn('w:val'), 'restart')
    tcPr.append(vm)

def _set_vmerge_continue(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:vMerge')):
        tcPr.remove(ex)
    vm = OxmlElement('w:vMerge')
    tcPr.append(vm)
    for child in list(tc):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            tc.remove(child)
    tc.append(OxmlElement('w:p'))

def _set_valign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(ex)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    tcPr.append(va)

def _set_row_keep_with_next(row):
    """Marca tots els paràgrafs de totes les cel·les d'una fila amb keep_with_next,
    perquè Word mai talli la pàgina entre aquesta fila i la següent."""
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True

def _remove_para_spacing(para):
    pPr = para._p.get_or_add_pPr()
    for ex in pPr.findall(qn('w:spacing')):
        pPr.remove(ex)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '0')
    pPr.append(sp)

def _cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
               color=None, red_suffix=None):
    """Escriu text a la primera fila d'una cel·la en Calibri.
    Si red_suffix no és None, afegeix un segon run en vermell (#FF0000)."""
    para = cell.paragraphs[0]
    para.alignment = align
    _remove_para_spacing(para)
    run = para.add_run(str(text))
    run.bold      = bold
    run.font.name = _FONT
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if red_suffix is not None:
        run_r = para.add_run(str(red_suffix))
        run_r.bold           = bold
        run_r.font.name      = _FONT
        run_r.font.size      = Pt(size)
        run_r.font.color.rgb = _COLOR_RED_RGB
    return run

def _write_dia_cell(cell, dia_setmana, data_text):
    """Escriu dia (negreta) i data (normal) en la cel·la DIA, centrats."""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_para_spacing(p1)
    r1 = p1.add_run(dia_setmana)
    r1.bold      = True
    r1.font.name = _FONT
    r1.font.size = Pt(9)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_para_spacing(p2)
    r2 = p2.add_run(data_text)
    r2.font.name = _FONT
    r2.font.size = Pt(9)


# ── Constructors de files de la taula ─────────────────────────────────────────

def _afegir_fila_tram(table, hora_txt, tema_txt, prof_txt, tema_color=None):
    """Afegeix una fila de tram (columnes HORA, TEMA, PROF; DIA s'omple per separat).
    tema_color: color del run TEMA (None = negre). Hora i professor sempre en negre."""
    row = table.add_row()
    for cell, txt, w, color in zip(
        row.cells[1:],
        [hora_txt,  tema_txt,   prof_txt],
        [_W_HORA,   _W_TEMA,    _W_PROF],
        [None,      tema_color, None],
    ):
        _set_cell_width(cell, w)
        _add_borders(cell)
        _cell_text(cell, txt, color=color)
    return row

def _afegir_fila_total(table, total_txt):
    """Fila de total d'hores: sombreig gris clar com a separador visual entre dies."""
    row = table.add_row()
    for cell, txt, w in zip(
        row.cells,
        [" ", " ", total_txt, " "],
        [_W_DIA, _W_HORA, _W_TEMA, _W_PROF],
    ):
        _set_cell_width(cell, w)
        _add_borders(cell)
        _set_shading(cell, _COLOR_TOTAL)
        _cell_text(cell, txt, bold=True)
    return row


# ── Generador principal ───────────────────────────────────────────────────────

def generar_document(horari_amb_professors, ruta_sortida,
                     estado_alumnos=None, estado_practicas=None,
                     estado_calendario=None, ruta_logo=RUTA_LOGO,
                     tipo_curso="mercancias", estado_franjas=None):
    """
    Genera el document .docx complet del CAP de mercaderies.
    Retorna ruta_sortida per a encadenament.
    """
    ampliacion        = _AMPLIACION_CURS.get(tipo_curso, set())
    estado_alumnos    = estado_alumnos    or {"alumnos": []}
    estado_practicas  = estado_practicas  or {"sesiones": []}
    estado_calendario = estado_calendario or {
        "fecha_inicio": {"conseguido": True, "valor": "—"},
        "dia_amarillo": {"conseguido": True, "valor": "—"},
        "festivos":     {"conseguido": True, "valor": []},
        "fecha_fin":    {"conseguido": True, "valor": "—"},
    }

    doc = Document()

    # ── Configuració de pàgina (A4, marges 0,5" lateral; 0,75" dalt per al header) ──
    section = doc.sections[0]
    section.page_width    = Twips(11906)
    section.page_height   = Twips(16838)
    section.left_margin   = Twips(720)
    section.right_margin  = Twips(720)
    section.top_margin    = Twips(1080)   # 0,75" per donar espai al header amb el logo
    section.bottom_margin = Twips(720)

    doc.styles['Normal'].font.name = _FONT
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

    # ── Header: logo petit a la dreta en totes les pàgines ───────────────────
    header = section.header
    hp     = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)
    hp.add_run().add_picture(ruta_logo, width=Inches(1.0))

    # ── PORTADA ───────────────────────────────────────────────────────────────

    # Logo gran centrat, amb espai al damunt
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(60)
    p_logo.paragraph_format.space_after  = Pt(55)
    p_logo.add_run().add_picture(ruta_logo, width=Inches(2.4))

    # Dades del curs: 6 línies negreta, centrades
    nom_curs = _NOM_CURS.get(tipo_curso, "MERCADERIES")
    texto_curs_portada = (
        _TEXT_CURS_PORTADA_VIATGERS if tipo_curso == "viatgers"
        else f"QUALIFICACIÓ INICIAL {nom_curs}"
    )
    for linia in [
        f"FORMACIÓ COMPLEMENTÀRIA: {DADES_CURS['formacio']}",
        f"CURS: {texto_curs_portada}",
        f"Durada: {DADES_CURS['durada']}",
        f"Empresa autoritzada: {DADES_CURS['empresa']}",
        f"NIF: {DADES_CURS['nif']}",
        f"Número de autorització: {DADES_CURS['autoritzacio']}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(linia)
        r.bold      = True
        r.font.name = _FONT
        r.font.size = Pt(12)

    # Títol (MERCADERIES o VIATGERS segons tipo_curso): gran, color de marca
    p_titol = doc.add_paragraph()
    p_titol.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titol.paragraph_format.space_before = Pt(80)
    p_titol.paragraph_format.space_after  = Pt(0)
    r = p_titol.add_run(nom_curs)
    r.bold           = True
    r.font.name      = _FONT
    r.font.size      = Pt(32)
    r.font.color.rgb = _COLOR_BRAND_RGB

    # Salt de pàgina
    p_salt = doc.add_paragraph()
    p_salt.paragraph_format.space_before = Pt(0)
    p_salt.paragraph_format.space_after  = Pt(0)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p_salt.add_run()._r.append(br)

    # ── TAULA: HORARI CLASSES TEÒRICA ─────────────────────────────────────────

    table = doc.add_table(rows=0, cols=4)
    _set_table_width(table, _W_TOTAL)
    _set_tbl_grid(table, [_W_DIA, _W_HORA, _W_TEMA, _W_PROF])
    _set_table_cell_margin(table, top=60, right=90, bottom=60, left=90)

    # Fila 1: "HORARI CLASSES TEÒRICA" (fusió horitzontal, fons de marca)
    fila_titol = table.add_row()
    cel_titol  = fila_titol.cells[0].merge(fila_titol.cells[3])
    _set_cell_width(cel_titol, _W_TOTAL)
    _add_borders(cel_titol)
    _set_shading(cel_titol, _COLOR_BRAND)
    _cell_text(cel_titol, "HORARI CLASSES TEÒRICA",
               bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
               color=_COLOR_WHITE_RGB)

    # Fila 2: capçaleres de columna (fons de marca, text blanc)
    fila_cap = table.add_row()
    for cell, txt, w in zip(
        fila_cap.cells,
        ["DIA", "HORA", "TEMA", "ESPECIALITAT PROFESSORAT"],
        [_W_DIA, _W_HORA, _W_TEMA, _W_PROF],
    ):
        _set_cell_width(cell, w)
        _add_borders(cell)
        _set_shading(cell, _COLOR_BRAND)
        _cell_text(cell, txt, bold=True, color=_COLOR_WHITE_RGB)

    # ── Dies i trams ──────────────────────────────────────────────────────────
    for entrada in horari_amb_professors:
        dia    = entrada["dia"]
        tramos = entrada["tramos"]

        dia_setmana = DIES_SETMANA_CA.get(dia.weekday(), "")
        data_text   = dia.strftime("%d/%m/%Y")

        files_dia = []
        for i, tramo in enumerate(tramos):
            hora_txt = _hora_rang(tramo["inicio"], tramo["fin"])

            if tramo.get("prueba_fuego"):
                # Prova de Foc: etiqueta específica, proveïdor com a professor, vermell
                tema_txt   = "MM.PP - Prova pràctica"
                prof_txt   = tramo.get("proveedor", tramo.get("profesor", ""))
                tema_color = _COLOR_RED_RGB
            elif tramo["tipo"] == "clase":
                tema_txt   = f"{tramo['codigo']} {tramo['nombre']}"
                prof_txt   = tramo.get("profesor", "")
                tema_color = _COLOR_RED_RGB if tramo["codigo"] in ampliacion else None
            elif "Transport" in tramo.get("nombre", ""):
                # Descansos de transport (anada/tornada): part del bloc PF, vermell
                tema_txt   = tramo["nombre"]
                prof_txt   = ""
                tema_color = _COLOR_RED_RGB
            else:
                # Descans normal de 15min
                tema_txt   = "DESCANS"
                prof_txt   = ""
                tema_color = None

            fila     = _afegir_fila_tram(table, hora_txt, tema_txt, prof_txt,
                                         tema_color=tema_color)
            files_dia.append(fila)
            dia_cell = fila.cells[0]
            _set_cell_width(dia_cell, _W_DIA)
            _add_borders(dia_cell)

            if i == 0:
                _set_vmerge_start(dia_cell)
                _set_valign(dia_cell, 'center')
                _write_dia_cell(dia_cell, dia_setmana, data_text)
            else:
                _set_vmerge_continue(dia_cell)

        fila_total = _afegir_fila_total(table, _total_hores_str(tramos))
        files_dia.append(fila_total)

        # Mantenir totes les files d'aquest dia juntes: cada fila (menys l'última,
        # la del total) s'"enganxa" a la següent amb keep_with_next -- si el dia
        # no cap sencer al que queda de pàgina, Word el salta sencer a la següent,
        # mai el parteix entre dues files. La fila del total NO s'enganxa a res,
        # perquè no s'ha de forçar que el dia següent quedi enganxat a aquest.
        for fila in files_dia[:-1]:
            _set_row_keep_with_next(fila)

    # ── PART DOS: alumnos, pràctiques, comunicacions ───────────────────────────

    def _titol_seccio(text, underline=True, space_before=22, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(8)
        r = p.add_run(text)
        r.bold           = True
        r.underline      = underline
        r.font.name      = _FONT
        r.font.size      = Pt(13)
        r.font.color.rgb = color if color is not None else _COLOR_BRAND_RGB

    def _linia(text, bold=False, size=11, indent=None, sb=2, sa=2, italic=False,
               color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if indent is not None:
            p.paragraph_format.left_indent = indent
        r = p.add_run(text)
        r.bold      = bold
        r.italic    = italic
        r.font.name = _FONT
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color

    # ── Blocs d'horari de teoria i ampliació ─────────────────────────────────
    if estado_franjas is not None:
        lun_jue = estado_franjas["horario_lun_jue"]["valor"]
        vdv     = estado_franjas["horario_viernes"]["valor"]
        sab     = estado_franjas["horario_sabado"]["valor"]

        _linia(
            f"HORARIS TEÒRICA: De Dilluns a Dijous de {lun_jue['inicio']}h a "
            f"{lun_jue['fin']}h. Divendres de {vdv['inicio']}h a {vdv['fin']}h. "
            f"Dissabte de {sab['inicio']}h a {sab['fin']}h.",
            bold=True, size=11, sb=12, sa=4,
        )

        dies_amp = sorted({
            entrada["dia"]
            for entrada in horari_amb_professors
            if entrada["dia"].weekday() in (4, 5)
            and any(t["tipo"] == "clase" and t["codigo"] in ampliacion
                    for t in entrada["tramos"])
        })

        if dies_amp:
            dates_txt = ", ".join(d.strftime("%d/%m/%Y") for d in dies_amp)
            _linia(
                f"AMPLIACIÓ (Cap 35 hores): Els alumnes d'ampliació assisteixen els "
                f"divendres i dissabtes següents: {dates_txt}. "
                f"Horari: divendres de {vdv['inicio']}h a {vdv['fin']}h "
                f"i dissabtes de {sab['inicio']}h a {sab['fin']}h.",
                bold=True, size=11, sb=4, sa=12, color=_COLOR_RED_RGB,
            )

    # Salt de pàgina
    p_salt2 = doc.add_paragraph()
    p_salt2.paragraph_format.space_before = Pt(0)
    p_salt2.paragraph_format.space_after  = Pt(0)
    br2 = OxmlElement('w:br')
    br2.set(qn('w:type'), 'page')
    p_salt2.add_run()._r.append(br2)

    # ── Línia obligatòria d'inspecció telemàtica ──────────────────────────────
    _linia("AQUESTES DADES SON A EFECTES D'INSPECCIONS TELEMÀTIQUES.",
           bold=True, size=11, sb=22, sa=6)

    # ── Secció: COMUNICACIONS INICIAL DEL CURS ────────────────────────────────
    _titol_seccio("COMUNICACIONS INICIAL DEL CURS", underline=False, space_before=4)

    data_inici = estado_calendario["fecha_inicio"]["valor"]
    data_final = estado_calendario["dia_amarillo"]["valor"]

    for label, valor in [
        ("Correu Electronic: ",    "inspeccionscap@autoescolaolivella.com"),
        ("Telèfon de Contacte: ",  "686329958"),
        ("",                       f"CURS QUALIFICACIÓ INICIAL DE {nom_curs}"),
        ("Data Inici: ",           data_inici),
        ("Data Final: ",           data_final),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        if label:
            r1 = p.add_run(label)
            r1.bold      = True
            r1.font.name = _FONT
            r1.font.size = Pt(11)
        r2 = p.add_run(valor)
        r2.bold      = not bool(label)
        r2.font.name = _FONT
        r2.font.size = Pt(11)

    # Salt de pàgina (separa COMUNICACIONS, que ara ve just després de l'horari,
    # de la secció d'alumnes/pràctiques que abans començava aquí mateix)
    p_salt_comm = doc.add_paragraph()
    p_salt_comm.paragraph_format.space_before = Pt(0)
    p_salt_comm.paragraph_format.space_after  = Pt(0)
    br_comm = OxmlElement('w:br')
    br_comm.set(qn('w:type'), 'page')
    p_salt_comm.add_run()._r.append(br_comm)

    # ── Secció 1: RELACIÓ ALUMNES 130H ───────────────────────────────────────
    _titol_seccio("RELACIÓ ALUMNES 130H", space_before=0)

    alumnes_complets = [a for a in estado_alumnos["alumnos"] if a["tipo_curso"] == "completo"]
    for i, alumne in enumerate(alumnes_complets, start=1):
        _linia(f"{i}. {alumne['nombre'].upper()}", size=11, sb=5, sa=0)
        _linia(alumne["documento"], size=11, indent=Inches(0.3), sb=0, sa=5)

    # ── Secció 2: AMPLIACIÓ ───────────────────────────────────────────────────
    _titol_seccio("AMPLIACIO:", underline=False, color=_COLOR_RED_RGB)

    alumnes_ampliacio = [a for a in estado_alumnos["alumnos"] if a["tipo_curso"] == "ampliacion"]
    if alumnes_ampliacio:
        for i, alumne in enumerate(alumnes_ampliacio, start=1):
            _linia(f"{i}. {alumne['nombre'].upper()}", size=11, sb=5, sa=0,
                   color=_COLOR_RED_RGB)
            _linia(alumne["documento"], size=11, indent=Inches(0.3), sb=0, sa=5,
                   color=_COLOR_RED_RGB)
    else:
        _linia("(No hi ha alumnes d'ampliació)", size=11)

    # ── Secció 3: HORARIS PRÀCTICA 130H ──────────────────────────────────────
    _titol_seccio("HORARIS PRÀCTICA 130H")

    # Professor: Rosa l'indica via estado_practicas; la resta el calcula el motor.
    professor = (estado_practicas or {}).get("profesor") or "—"

    alumnes_c = sum(1 for a in estado_alumnos["alumnos"] if a["tipo_curso"] == "completo")
    alumnes_a = sum(1 for a in estado_alumnos["alumnos"] if a["tipo_curso"] == "ampliacion")
    dies_lect  = construir_dias_lectivos(estado_calendario)
    franges_p  = construir_franjas_semanales(estado_franjas)
    sessions_calc = colocar_practicas(dies_lect, franges_p, alumnes_c, alumnes_a)

    if sessions_calc:
        for sessio in sessions_calc:
            inici = _fmt_hora_practica(sessio["hora_inicio"].strftime("%H:%M"))
            fi    = _fmt_hora_practica(sessio["hora_fin"].strftime("%H:%M"))
            color = _COLOR_RED_RGB if sessio["ampliacio"] else None
            _linia(
                f"{sessio['data'].strftime('%d/%m/%Y')}  {inici} a {fi}  "
                f"Professor: {professor}",
                size=11, sb=3, sa=3, color=color,
            )
    else:
        _linia("(No hi ha sessions de pràctiques registrades)", size=11)

    # ── PART TRES: CRONOGRAMA PER SETMANES (pàgina apaïsada) ─────────────────
    cronograma = cronograma_por_semanas(horari_amb_professors)

    # Selecció de la taula de matèries i categories segons el tipus de curs
    mat_cron, cat_cron = _MATERIAS_CURS.get(tipo_curso, _MATERIAS_CURS["mercancias"])
    n_mat = len(mat_cron)          # 16 per mercancías, 19 per viatgers
    col_total_idx = n_mat + 1      # índex de la columna "TOTAL" (última)

    seccio_cron = doc.add_section(WD_SECTION.NEW_PAGE)
    seccio_cron.page_width    = Twips(16838)
    seccio_cron.page_height   = Twips(11906)
    seccio_cron.left_margin   = Twips(360)
    seccio_cron.right_margin  = Twips(360)
    seccio_cron.top_margin    = Twips(720)
    seccio_cron.bottom_margin = Twips(540)
    pgSz = seccio_cron._sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.set(qn('w:orient'), 'landscape')

    # Títol del cronograma
    p_tit_cron = doc.add_paragraph()
    p_tit_cron.paragraph_format.space_before = Pt(0)
    p_tit_cron.paragraph_format.space_after  = Pt(10)
    r = p_tit_cron.add_run("CRONOGRAMA D'HORES PER SETMANA")
    r.bold           = True
    r.font.name      = _FONT
    r.font.size      = Pt(13)
    r.font.color.rgb = _COLOR_BRAND_RGB

    # Amplada de cada columna de matèria (dinàmica: s'ajusta per a 16 o 19 columnes)
    wc_mat     = (16118 - _WC_SETMANA - _WC_TOTAL) // n_mat
    col_widths = [_WC_SETMANA] + [wc_mat] * n_mat + [_WC_TOTAL]
    w_taula    = sum(col_widths)

    taula_cron = doc.add_table(rows=0, cols=n_mat + 2)
    _set_table_width(taula_cron, w_taula)
    _set_tbl_grid(taula_cron, col_widths)
    _set_table_cell_margin(taula_cron, top=40, right=50, bottom=40, left=50)

    # ── Fila 0: categories (fusions horitzontals, fons de marca) ──────────────
    fila_cat = taula_cron.add_row()

    c0 = fila_cat.cells[0]
    _set_cell_width(c0, col_widths[0])
    _add_borders(c0)
    _set_shading(c0, _COLOR_BRAND)

    for label, idx_ini, idx_fi in cat_cron:
        col_ini = idx_ini + 1
        col_fi  = idx_fi  + 1
        merged  = fila_cat.cells[col_ini].merge(fila_cat.cells[col_fi])
        _set_cell_width(merged, wc_mat * (idx_fi - idx_ini + 1))
        _add_borders(merged)
        _set_shading(merged, _COLOR_BRAND)
        _cell_text(merged, label, bold=True, size=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=_COLOR_WHITE_RGB)

    c_tot_cat = fila_cat.cells[col_total_idx]
    _set_cell_width(c_tot_cat, col_widths[col_total_idx])
    _add_borders(c_tot_cat)
    _set_shading(c_tot_cat, _COLOR_BRAND)

    # ── Fila 1: capçaleres de columna (fons de marca, text blanc) ─────────────
    fila_cap = taula_cron.add_row()

    c = fila_cap.cells[0]
    _set_cell_width(c, col_widths[0])
    _add_borders(c)
    _set_shading(c, _COLOR_BRAND)
    _cell_text(c, "Setmana", bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER, color=_COLOR_WHITE_RGB)

    for j, (codigo, nom_curt) in enumerate(mat_cron):
        c = fila_cap.cells[j + 1]
        _set_cell_width(c, col_widths[j + 1])
        _add_borders(c)
        _set_shading(c, _COLOR_BRAND)
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _remove_para_spacing(p1)
        r1 = p1.add_run(codigo)
        r1.bold           = True
        r1.font.name      = _FONT
        r1.font.size      = Pt(8)
        r1.font.color.rgb = _COLOR_WHITE_RGB
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _remove_para_spacing(p2)
        r2 = p2.add_run(nom_curt)
        r2.font.name      = _FONT
        r2.font.size      = Pt(7)
        r2.font.color.rgb = _COLOR_WHITE_RGB

    c = fila_cap.cells[col_total_idx]
    _set_cell_width(c, col_widths[col_total_idx])
    _add_borders(c)
    _set_shading(c, _COLOR_BRAND)
    for txt in ["TOTAL", "HORES SET."]:
        p = c.paragraphs[0] if txt == "TOTAL" else c.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _remove_para_spacing(p)
        r = p.add_run(txt)
        r.bold           = True
        r.font.name      = _FONT
        r.font.size      = Pt(8)
        r.font.color.rgb = _COLOR_WHITE_RGB

    # ── Files de dades (una per setmana) ──────────────────────────────────────
    for num_sem, materias in sorted(cronograma.items()):
        fila = taula_cron.add_row()

        c = fila.cells[0]
        _set_cell_width(c, col_widths[0])
        _add_borders(c)
        _cell_text(c, f"Setmana {num_sem}", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        total_sem = 0.0
        for j, (codigo, _) in enumerate(mat_cron):
            c = fila.cells[j + 1]
            _set_cell_width(c, col_widths[j + 1])
            _add_borders(c)
            h = materias.get(codigo, 0.0)
            total_sem += h
            if h and codigo in ampliacion:
                txt = _fmt_hores_cron(h)
                _cell_text(c, txt + "/", size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
                           red_suffix=txt)
            else:
                _cell_text(c, _fmt_hores_cron(h) if h else "", size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER)

        c = fila.cells[col_total_idx]
        _set_cell_width(c, col_widths[col_total_idx])
        _add_borders(c)
        _cell_text(c, _fmt_hores_cron(total_sem), bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Fila de totals per matèria ─────────────────────────────────────────────
    fila_tot = taula_cron.add_row()

    c = fila_tot.cells[0]
    _set_cell_width(c, col_widths[0])
    _add_borders(c)
    _set_shading(c, _COLOR_TOTAL)
    _cell_text(c, "TOTAL", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for j, (codigo, _) in enumerate(mat_cron):
        c = fila_tot.cells[j + 1]
        _set_cell_width(c, col_widths[j + 1])
        _add_borders(c)
        _set_shading(c, _COLOR_TOTAL)
        h_tot = sum(sem.get(codigo, 0.0) for sem in cronograma.values())
        txt = _fmt_hores_cron(h_tot)
        if codigo in ampliacion:
            _cell_text(c, txt + "/", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER, red_suffix=txt)
        else:
            _cell_text(c, txt, bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    c = fila_tot.cells[col_total_idx]
    _set_cell_width(c, col_widths[col_total_idx])
    _add_borders(c)
    _set_shading(c, _COLOR_TOTAL)
    _cell_text(c, "130 h", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(ruta_sortida)
    return ruta_sortida


# ── Script directe ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    from datetime import date, time as time_
    from datos_ejemplo import estados_ejemplo
    from ensamblaje import generar_horario, aplicar_profesores
    from herramientas.motor_prueba_fuego import crear_prueba_fuego

    estados = estados_ejemplo()
    pf = crear_prueba_fuego(
        fecha=date(2026, 3, 21),
        hora_inicio=time_(10, 0),
    )
    resultat = generar_horario(
        estados["calendario"],
        estados["franjas"],
        estados["orden"],
        tipo_curso="mercancias",
        prueba_fuego=pf,
    )
    horari = aplicar_profesores(resultat["horario"], estados["profesores"])

    ruta = generar_document(
        horari, str(_BASE / "output_cap.docx"),
        estado_alumnos=estados["alumnos"],
        estado_practicas=estados["practicas"],
        estado_calendario=estados["calendario"],
        estado_franjas=estados["franjas"],
        tipo_curso="mercancias",
    )
    print(f"Document generat: {ruta}")
